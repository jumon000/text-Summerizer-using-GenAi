import PyPDF2
import docx

def extract_text_from_pdf(pdf_path):
    """Extracts text from a PDF file."""
    text = ""
    with open(pdf_path, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            if page.extract_text():
                text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(docx_path):
    """Extracts text from a DOCX file, preserving formatting like lists and tables."""
    doc = docx.Document(docx_path)
    extracted_text = []

    for para in doc.paragraphs:
        extracted_text.append(para.text.strip())

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            extracted_text.append(" | ".join(row_text))  # Formatting table rows

    return "\n".join([line for line in extracted_text if line])

def save_extracted_text(file_path, output_file="extracted_text.txt"):
    """Extracts text from PDF or DOCX and saves it to a text file."""
    if file_path.endswith(".pdf"):
        text = extract_text_from_pdf(file_path)
    elif file_path.endswith(".docx"):
        text = extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format! Please provide a PDF or DOCX.")

    with open(output_file, "w", encoding="utf-8") as file:
        file.write(text)

def preprocess_text(text):
    """
    Preprocesses the extracted text by:
    - Removing extra whitespace and newlines
    - Retaining essential punctuation
    - Handling bullet points and numbered lists
    - Normalizing text (e.g., lowercasing, if needed)
    """
    # Normalize whitespace and remove extra newlines
    text = re.sub(r"\s+", " ", text).strip()

    # Retain only useful characters while preserving punctuation
    text = re.sub(r"[^\w\s.,;:!?()\-*/%&]", "", text)

    # Normalize text (optional: convert to lowercase)
    text = text.lower()

    return text

